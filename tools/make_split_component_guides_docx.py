from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DOWNLOADS = Path(r"C:\downloads")


FILES = {
    "vcl_user": "VCL_CollapsibleDBGrid_User_Guide.docx",
    "vcl_engineering": "VCL_CollapsibleDBGrid_Engineering_Guide.docx",
    "fmx_user": "FMX_CollapsibleDBGrid_User_Guide.docx",
    "fmx_engineering": "FMX_CollapsibleDBGrid_Engineering_Guide.docx",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Pt(width / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def apply_styles(doc, title):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(0x2E, 0x74, 0xB5), 18, 10),
        ("Heading 2", 13, RGBColor(0x2E, 0x74, 0xB5), 14, 7),
        ("Heading 3", 12, RGBColor(0x1F, 0x4D, 0x78), 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title_style = styles.add_style("Guide Title", 1)
    title_style.font.name = "Calibri"
    title_style.font.size = Pt(20)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    title_style.paragraph_format.space_after = Pt(3)

    code_style = styles.add_style("Code Block", 1)
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(9)
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(3)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(title)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph(style="Guide Title")
    p.add_run(title)
    sub = doc.add_paragraph()
    run = sub.add_run(subtitle)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_code(doc, text):
    for line in text.strip("\n").splitlines():
        p = doc.add_paragraph(style="Code Block")
        p.add_run(line)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F4F6F9")
        p._p.get_or_add_pPr().append(shd)


def add_table(doc, rows, widths=None):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    if widths is None:
        widths = [2600, 6760] if len(rows[0]) == 2 else [9360 // len(rows[0])] * len(rows[0])
    set_table_width(table, widths)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(value)
            if r_idx == 0:
                run.bold = True
                set_cell_shading(cell, "E8EEF5")
    doc.add_paragraph()


def build_common_user_sections(doc, component_name):
    doc.add_paragraph("Required Data Components", style="Heading 1")
    add_bullets(doc, [
        "A database connection component, such as TFDConnection.",
        "A dataset component, such as TFDQuery, TFDTable, or another TDataSet descendant.",
        "A TDataSource connected to the dataset.",
        f"A {component_name} dropped on the form at design time.",
    ])
    add_code(doc, """
FDQuery1.Connection := FDConnection1;
DataSource1.DataSet := FDQuery1;
CollapsibleDBGrid1.DataSource := DataSource1;
""")

    doc.add_paragraph("Basic Setup", style="Heading 1")
    add_steps(doc, [
        "Drop the component on the form.",
        "Set DataSource.",
        "Set GroupField to the dataset field used for grouping.",
        "Add items to Columns.",
        "For each column, set FieldName, Header, and Width.",
        "Optionally set KeyField to a stable unique field.",
    ])

    doc.add_paragraph("Core Properties", style="Heading 1")
    add_table(doc, [
        ["Property", "Purpose"],
        ["DataSource", "TDataSource used by the grid."],
        ["GroupField", "Dataset field used to create group header rows."],
        ["KeyField", "Optional stable row identity field."],
        ["StartCollapsed", "Starts groups collapsed by default."],
        ["AutoExpandSelectedGroup", "Expands the group containing the selected dataset row."],
        ["AutoHeight", "Resizes the grid vertically to full visible rows where possible."],
        ["AutoSizeColumnsAtStartup", "When true, autosizes columns once at runtime startup. Default is false so collection Width values are authoritative."],
        ["ShowDesignPreview", "Shows dummy preview rows at design time when no active dataset is available."],
        ["UseActiveDataAtDesignTime", "Allows an already-active dataset to be displayed in the designer."],
        ["DesignPreviewGroupCount", "Number of preview groups."],
        ["DesignPreviewRowsPerGroup", "Number of preview rows per preview group."],
        ["ShowEmptyGroupAs", "Text used for blank or null group values."],
        ["RemoveFileNamePath", "Displays path-like values without directory prefixes."],
        ["Columns", "Collection of displayed fields."],
    ])

    doc.add_paragraph("Column Items", style="Heading 1")
    add_table(doc, [
        ["Property", "Purpose"],
        ["FieldName", "Dataset field displayed in the column."],
        ["Header", "Text shown in the header row."],
        ["Width", "Design width and fallback runtime width."],
    ])

    doc.add_paragraph("Events", style="Heading 1")
    add_table(doc, [
        ["Event", "Purpose"],
        ["OnGetGroupText", "Customize group header text."],
        ["OnGetCellText", "Customize cell text."],
        ["OnBeforeGroupExpand", "Allows group expansion to be approved or cancelled."],
        ["OnBeforeGroupCollapse", "Allows group collapse to be approved or cancelled."],
        ["OnBeforeRowSelect", "Allows dataset row selection to be approved or cancelled."],
        ["OnFormatGroupHeader", "Customizes final group header display text."],
        ["OnDrawGroupRow", "Allows custom drawing for group rows."],
        ["OnDrawDataRow", "Allows custom drawing for data rows."],
        ["OnGroupExpanded", "Runs after a group opens."],
        ["OnGroupCollapsed", "Runs after a group closes."],
        ["OnRowActivated", "Runs when a data row is activated."],
        ["OnSelectionChanged", "Runs when selected row changes."],
    ])

    doc.add_paragraph("User Interaction", style="Heading 1")
    add_bullets(doc, [
        "Click a group row to expand or collapse it.",
        "Expanded groups show [-]; collapsed groups show [+].",
        "Right-click the grid for Expand All Groups and Collapse All Groups.",
        "Clicking a data row moves the dataset to that row when synchronization is possible.",
        "Design-time preview rows use configured column Header or FieldName values when no dataset is active.",
        "When AutoHeight is enabled, allow enough form or parent-control space for larger row fonts and expanded groups.",
    ])


def build_vcl_user(path):
    doc = Document()
    apply_styles(doc, "VCL CollapsibleDBGrid User Guide")
    add_title(doc, "VCL CollapsibleDBGrid User Guide", "User setup and operation guide for TVCLCollapsibleDBGrid.")
    doc.add_paragraph("Overview", style="Heading 1")
    doc.add_paragraph("TVCLCollapsibleDBGrid is a design-time, data-aware, read-only grouped grid for VCL forms. It inherits from TCustomDrawGrid and renders headers, group rows, data rows, selection, and alternating row colors with VCL Canvas drawing.")
    build_common_user_sections(doc, "TVCLCollapsibleDBGrid")
    doc.add_paragraph("VCL Styling Properties", style="Heading 1")
    add_table(doc, [
        ["Property", "Type / Purpose"],
        ["Font", "TFont for normal data rows."],
        ["HeaderFont", "TFont for column header text."],
        ["HeaderRowColor", "TColor for column header background."],
        ["GroupFont", "TFont for group header rows."],
        ["GroupRowColor", "TColor for group row background."],
        ["OddRowColor", "TColor for alternating odd data rows."],
        ["SelectedRowColor", "TColor for selected data rows."],
        ["Color", "Base background color for normal even rows."],
        ["ScrollBars", "Inherited TScrollStyle exposed for controlling VCL grid scroll bar behavior."],
    ])
    doc.add_paragraph("VCL Install", style="Heading 1")
    add_steps(doc, [
        "Close Delphi and verify bds.exe is not running.",
        "Build CommonCollapsibleDBGridRuntime.dpk.",
        "Build VCLCollapsibleDBGridRuntime.dpk.",
        "Build VCLCollapsibleDBGridDesign.dpk.",
        "Install VCLCollapsibleDBGridDesign.dpk.",
        "Restart Delphi.",
        "Confirm TVCLCollapsibleDBGrid appears on the CollapsibleDBGrid palette page.",
    ])
    doc.add_paragraph("VCL Notes", style="Heading 1")
    add_bullets(doc, [
        "Do not expect the component to open database connections automatically in the IDE.",
        "Use the Columns collection, not runtime-created grid columns.",
        "HeaderFont and GroupFont are normal VCL TFont objects and use the native font editor.",
        "Changing Font, HeaderFont, or GroupFont recalculates row metrics and refreshes the VCL scroll range.",
        "Large row fonts may require a taller grid or parent form when AutoHeight is enabled.",
    ])
    doc.save(path)


def build_fmx_user(path):
    doc = Document()
    apply_styles(doc, "FMX CollapsibleDBGrid User Guide")
    add_title(doc, "FMX CollapsibleDBGrid User Guide", "User setup and operation guide for TFMXCollapsibleDBGrid.")
    doc.add_paragraph("Overview", style="Heading 1")
    doc.add_paragraph("TFMXCollapsibleDBGrid is a design-time, data-aware, read-only grouped grid for FMX forms. It inherits from TStringGrid but uses its own Columns collection as the public data-column model.")
    build_common_user_sections(doc, "TFMXCollapsibleDBGrid")
    doc.add_paragraph("FMX Styling Properties", style="Heading 1")
    add_table(doc, [
        ["Property", "Type / Purpose"],
        ["TextSettings", "Inherited FMX text settings for normal data rows."],
        ["HeaderFontFamily", "TFontName for column header font family."],
        ["HeaderFontPointSize", "Integer point size for column header font."],
        ["HeaderFontColor", "TAlphaColor for column header text."],
        ["HeaderFontStyle", "TFontStyles for column header style."],
        ["HeaderRowColor", "TAlphaColor for column header background."],
        ["GroupFontFamily", "TFontName for group row font family."],
        ["GroupFontPointSize", "Integer point size for group row font."],
        ["GroupFontColor", "TAlphaColor for group row text."],
        ["GroupFontStyle", "TFontStyles for group row style."],
        ["GroupRowColor", "TAlphaColor for group row background."],
        ["OddRowColor", "TAlphaColor for alternating odd data rows."],
        ["SelectedRowColor", "TAlphaColor for selected data rows."],
    ])
    doc.add_paragraph("FMX Install", style="Heading 1")
    add_steps(doc, [
        "Close Delphi and verify bds.exe is not running.",
        "Build CommonCollapsibleDBGridRuntime.dpk.",
        "Build FMXCollapsibleDBGridRuntime.dpk.",
        "Build FMXCollapsibleDBGridDesign.dpk.",
        "Install FMXCollapsibleDBGridDesign.dpk.",
        "Restart Delphi.",
        "Confirm TFMXCollapsibleDBGrid appears on the CollapsibleDBGrid palette page.",
    ])
    doc.add_paragraph("FMX Notes", style="Heading 1")
    add_bullets(doc, [
        "Do not add native TStringColumn children manually.",
        "Use the component Columns collection for data columns.",
        "HeaderFontFamily and GroupFontFamily use TFontName so the Object Inspector can show installed fonts.",
        "HeaderFontPointSize and GroupFontPointSize are Integer properties to avoid FMX animation editors.",
        "FMX visual columns are internal and are marked not stored.",
    ])
    doc.save(path)


def build_common_engineering_sections(doc, framework):
    doc.add_paragraph("Shared Core", style="Heading 1")
    add_table(doc, [
        ["Unit", "Responsibility"],
        ["CollapsibleDBGrid.Core.pas", "Row kinds, row records, row list, and collapse state."],
        ["CollapsibleDBGrid.Data.pas", "Column collection, build options, dataset projection, design preview rows."],
    ])
    doc.add_paragraph("Data Projection", style="Heading 1")
    add_bullets(doc, [
        "BuildFromDataSet converts dataset records into visible group and data rows.",
        "GroupField chooses the grouping value.",
        "Columns chooses displayed fields.",
        "GroupState decides whether data rows are included for each group.",
        "BuildPreviewRows creates design-time dummy groups and rows.",
        "Preview rows use each column Header, then FieldName, then a generic column caption.",
        "RemoveFileNamePath strips path prefixes from path-like displayed values.",
    ])
    doc.add_paragraph("Dataset Synchronization", style="Heading 1")
    add_bullets(doc, [
        "Each framework component owns a TDataLink descendant.",
        "ActiveChanged and DataSetChanged refresh visible rows.",
        "DataSetScrolled synchronizes grid selection from the dataset.",
        "Clicking a data row attempts to move the dataset to that row.",
        "OnBeforeRowSelect can cancel row-to-dataset synchronization.",
        "Current synchronization primarily uses RecNo; KeyField is captured for future stronger locate behavior.",
    ])
    doc.add_paragraph("Package Order", style="Heading 1")
    add_steps(doc, [
        "Build CommonCollapsibleDBGridRuntime.dpk.",
        f"Build {framework}CollapsibleDBGridRuntime.dpk.",
        f"Build {framework}CollapsibleDBGridDesign.dpk.",
        "Install only the design package.",
    ])


def add_ide_install_tutorial(doc, framework):
    runtime_package = f"{framework}CollapsibleDBGridRuntime.dpk"
    design_package = f"{framework}CollapsibleDBGridDesign.dpk"
    component_name = f"T{framework}CollapsibleDBGrid"
    doc.add_paragraph("IDE Installation Tutorial", style="Heading 1")
    doc.add_paragraph(
        "Use this process when installing the component into Delphi manually from the package files. "
        "Install the design-time package only; the design package depends on the runtime packages."
    )
    doc.add_paragraph("Before Installing", style="Heading 2")
    add_steps(doc, [
        "Close all running projects in Delphi.",
        "Close the Delphi IDE completely.",
        "Open Task Manager and confirm bds.exe is not running.",
        "Start Delphi again with no project loaded.",
        "Choose File > Open Project.",
        "Open packages\\CommonCollapsibleDBGridRuntime.dpk.",
        "Right-click the package in Project Manager and choose Build.",
        f"Choose File > Open Project and open packages\\{runtime_package}.",
        "Right-click the runtime package in Project Manager and choose Build.",
        f"Choose File > Open Project and open packages\\{design_package}.",
        "Right-click the design package in Project Manager and choose Build.",
        "Right-click the design package again and choose Install.",
        "Confirm Delphi reports that the component was installed.",
    ])
    doc.add_paragraph("Verify the Install", style="Heading 2")
    add_steps(doc, [
        "Create a new test project for the matching framework.",
        "Open the Tool Palette.",
        "Search for CollapsibleDBGrid.",
        f"Confirm {component_name} appears on the palette.",
        "Drop the component on a blank design-time form.",
        "Confirm the Columns collection is visible in the Object Inspector.",
        "Add at least one column and confirm the design preview appears.",
    ])
    doc.add_paragraph("Project Search Path", style="Heading 2")
    add_steps(doc, [
        "Open Tools > Options > Language > Delphi > Library.",
        "Select the Win32 platform unless your package build targets a different platform.",
        "Add the CollapsibleDBGrid source folder to the Library path if it is not already present.",
        "Include src\\Common and the matching framework source folder.",
        "Restart Delphi after changing the library path.",
    ])
    doc.add_paragraph("If Delphi Refuses to Load the Package", style="Heading 2")
    add_bullets(doc, [
        "Close Delphi and verify bds.exe is not running before rebuilding or replacing BPL files.",
        "Make sure CommonCollapsibleDBGridRuntime is built before the framework runtime package.",
        "Make sure the framework runtime package is built before the design package.",
        "Do not install runtime packages manually; install only the design package.",
        "If a duplicate-unit package conflict appears, uninstall the older conflicting design package and restart Delphi.",
        "If the component is missing from the palette, rebuild and reinstall the design package, then restart Delphi.",
    ])


def build_vcl_engineering(path):
    doc = Document()
    apply_styles(doc, "VCL CollapsibleDBGrid Engineering Guide")
    add_title(doc, "VCL CollapsibleDBGrid Engineering Guide", "Implementation and maintenance guide for TVCLCollapsibleDBGrid.")
    doc.add_paragraph("Component", style="Heading 1")
    add_code(doc, "TVCLCollapsibleDBGrid = class(TCustomDrawGrid, ICollapsibleDBGridColumnsChanged)")
    build_common_engineering_sections(doc, "VCL")
    add_ide_install_tutorial(doc, "VCL")
    doc.add_paragraph("VCL Rendering", style="Heading 1")
    add_bullets(doc, [
        "DrawCell paints fixed header row, group rows, selected rows, odd rows, and normal rows.",
        "FixedRows is 1 when Columns is populated.",
        "DataRowIndex subtracts FixedRows to map grid rows to visible projection rows.",
        "HeaderFont, GroupFont, HeaderRowColor, GroupRowColor, OddRowColor, SelectedRowColor, Color, and Font drive rendering.",
        "OnDrawGroupRow and OnDrawDataRow can replace default row drawing when DefaultDraw is set false.",
        "OnFormatGroupHeader can modify group header text before it is drawn.",
        "UpdateMetricsForFont recalculates row heights from Font, GroupFont, and HeaderFont.",
        "RefreshScrollBars posts a WM_SIZE-style recalculation so VCL scroll bars match the recalculated row heights at startup and after font changes.",
    ])
    doc.add_paragraph("VCL Design-Time Behavior", style="Heading 1")
    add_bullets(doc, [
        "ShowDesignPreview is true by default.",
        "UseActiveDataAtDesignTime is false by default.",
        "The component does not auto-open datasets or connections in the IDE.",
        "GroupFont and HeaderFont are owned TFont instances with OnChange handlers.",
        "ScrollBars is published so users can choose standard VCL scroll bar behavior from the Object Inspector.",
    ])
    doc.add_paragraph("VCL Maintenance Checklist", style="Heading 1")
    add_bullets(doc, [
        "Do not subclass TDBGrid; this component owns a projection model and custom drawing.",
        "Keep package dependencies separated between runtime and design-time packages.",
        "Treat compiler hints and warnings as important.",
        "When adding a design-time property, use Object Inspector-friendly VCL types.",
        "When row-height logic changes, verify startup rendering before and after the first grid interaction; stale scroll ranges can appear only on first show.",
        "When changing packages, close Delphi and verify bds.exe is not running before installing.",
    ])
    doc.save(path)


def build_fmx_engineering(path):
    doc = Document()
    apply_styles(doc, "FMX CollapsibleDBGrid Engineering Guide")
    add_title(doc, "FMX CollapsibleDBGrid Engineering Guide", "Implementation and maintenance guide for TFMXCollapsibleDBGrid.")
    doc.add_paragraph("Component", style="Heading 1")
    add_code(doc, "TFMXCollapsibleDBGrid = class(TStringGrid, ICollapsibleDBGridColumnsChanged)")
    build_common_engineering_sections(doc, "FMX")
    add_ide_install_tutorial(doc, "FMX")
    doc.add_paragraph("FMX Rendering", style="Heading 1")
    add_bullets(doc, [
        "DefaultDrawing is false.",
        "OnDrawColumnBackground paints row backgrounds.",
        "OnDrawColumnCell paints cell text.",
        "OnDrawColumnHeader paints the header row.",
        "RowFillColor chooses GroupRowColor, SelectedRowColor, OddRowColor, or normal white.",
        "OnDrawGroupRow and OnDrawDataRow can replace default row drawing when DefaultDraw is set false.",
        "OnFormatGroupHeader can modify group header text before it is drawn.",
        "Header style uses HeaderFontFamily, HeaderFontPointSize, HeaderFontColor, HeaderFontStyle, and HeaderRowColor.",
        "Group style uses GroupFontFamily, GroupFontPointSize, GroupFontColor, GroupFontStyle, and GroupRowColor.",
    ])
    doc.add_paragraph("FMX Object Inspector Rules", style="Heading 1")
    add_table(doc, [
        ["Rule", "Reason"],
        ["Use TFontName for font family properties.", "Plain string does not show the installed-font editor."],
        ["Use Integer point-size properties.", "Single can trigger FMX animation menu choices."],
        ["Use TAlphaColor for color properties.", "Matches FMX color model."],
        ["Avoid publishing raw TTextSettings.", "Its editable members are public, not published, so the Object Inspector shows no useful child properties."],
        ["Keep internal TStringColumn objects not stored.", "Prevents FMX forms from streaming visual columns separately from data columns."],
    ])
    doc.add_paragraph("FMX Design-Time Behavior", style="Heading 1")
    add_bullets(doc, [
        "ShowDesignPreview is true by default.",
        "UseActiveDataAtDesignTime is false by default.",
        "The component does not auto-open datasets or connections in the IDE.",
        "DefinePresentationName uses the TStringGrid presentation so FMX styled presentation is registered.",
        "RebuildColumns creates internal TStringColumn instances from the public Columns collection.",
    ])
    doc.add_paragraph("FMX Maintenance Checklist", style="Heading 1")
    add_bullets(doc, [
        "Do not expose FMX properties using types that create unwanted animation editors.",
        "Do not publish raw settings objects unless their child properties are published and editable.",
        "Keep native FMX visual columns internal.",
        "Verify header and group styling in the Object Inspector after changing published property types.",
        "When changing packages, close Delphi and verify bds.exe is not running before installing.",
    ])
    doc.save(path)


def main():
    DOCS.mkdir(parents=True, exist_ok=True)
    DOWNLOADS.mkdir(parents=True, exist_ok=True)

    builders = {
        "vcl_user": build_vcl_user,
        "vcl_engineering": build_vcl_engineering,
        "fmx_user": build_fmx_user,
        "fmx_engineering": build_fmx_engineering,
    }

    for key, builder in builders.items():
        for folder in (DOCS, DOWNLOADS):
            builder(folder / FILES[key])
            print(folder / FILES[key])


if __name__ == "__main__":
    main()
