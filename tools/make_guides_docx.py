from pathlib import Path
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DOWNLOADS = Path(r"C:\downloads")


PRESET = {
    "font": "Calibri",
    "body_size": 11,
    "h1_size": 16,
    "h2_size": 13,
    "h3_size": 12,
    "heading_blue": RGBColor(0x2E, 0x74, 0xB5),
    "heading_dark": RGBColor(0x1F, 0x4D, 0x78),
    "header_fill": "E8EEF5",
    "code_fill": "F4F6F9",
    "border": "D9E2F3",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.allow_autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Pt(width / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def style_document(doc, title):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = PRESET["font"]
    normal.font.size = Pt(PRESET["body_size"])
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, PRESET["heading_blue"], 18, 10),
        ("Heading 2", 13, PRESET["heading_blue"], 14, 7),
        ("Heading 3", 12, PRESET["heading_dark"], 10, 5),
    ]:
        st = styles[name]
        st.font.name = PRESET["font"]
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    title_style = styles.add_style("Guide Title", 1)
    title_style.font.name = PRESET["font"]
    title_style.font.size = Pt(20)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    title_style.paragraph_format.space_after = Pt(3)

    code_style = styles.add_style("Code Block", 1)
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(9)
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(3)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = title
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_code_paragraph(doc, text):
    p = doc.add_paragraph(style="Code Block")
    run = p.add_run(text if text else " ")
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), PRESET["code_fill"])
    p_pr.append(shd)
    return p


def add_inline_runs(paragraph, text):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def add_simple_table(doc, rows):
    if not rows:
        return
    col_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    widths = [2800, 6560] if col_count == 2 else [9360 // col_count] * col_count
    set_table_width(table, widths)
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline_runs(p, value.strip())
            if r == 0:
                set_cell_shading(cell, PRESET["header_fill"])
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()


def markdown_to_docx(md_path, docx_path, subtitle):
    title = md_path.stem.replace("_", " ")
    doc = Document()
    style_document(doc, title)

    p = doc.add_paragraph(style="Guide Title")
    p.add_run(title.replace("CollapsibleDBGrid ", "CollapsibleDBGrid: "))
    sub = doc.add_paragraph()
    run = sub.add_run(subtitle)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    in_code = False
    code_lines = []
    table_rows = []

    def flush_code():
        nonlocal code_lines
        if code_lines:
            for line in code_lines:
                add_code_paragraph(doc, line)
            code_lines = []

    def flush_table():
        nonlocal table_rows
        if table_rows:
            cleaned = []
            for row in table_rows:
                cells = [c.strip() for c in row.strip().strip("|").split("|")]
                if all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
                    continue
                cleaned.append(cells)
            add_simple_table(doc, cleaned)
            table_rows = []

    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_rows.append(line)
            continue
        flush_table()
        if not line.strip():
            continue
        if line.startswith("# "):
            # Title already handled.
            continue
        if line.startswith("## "):
            flush_code()
            doc.add_paragraph(line[3:], style="Heading 1")
        elif line.startswith("### "):
            flush_code()
            doc.add_paragraph(line[4:], style="Heading 2")
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, line[2:])
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, re.sub(r"^\d+\. ", "", line))
        else:
            p = doc.add_paragraph()
            add_inline_runs(p, line)

    flush_code()
    flush_table()
    doc.core_properties.title = title
    doc.core_properties.subject = subtitle
    doc.save(docx_path)


def main():
    DOWNLOADS.mkdir(parents=True, exist_ok=True)
    outputs = [
        (
            DOCS / "CollapsibleDBGrid_User_Guide.md",
            DOWNLOADS / "CollapsibleDBGrid_User_Guide.docx",
            "User setup and operation guide for both VCL and FMX components.",
        ),
        (
            DOCS / "CollapsibleDBGrid_Engineering_Guide.md",
            DOWNLOADS / "CollapsibleDBGrid_Engineering_Guide.docx",
            "Engineering and maintenance guide for the shared core, VCL, and FMX implementations.",
        ),
    ]
    for src, dst, subtitle in outputs:
        markdown_to_docx(src, dst, subtitle)
        print(dst)


if __name__ == "__main__":
    sys.exit(main())
